from pathlib import Path
import csv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OLD = ROOT / "revision_outputs"
NEW = ROOT / "strengthened_revision_outputs"
FINAL = NEW / "final_manuscript"
FINAL.mkdir(parents=True, exist_ok=True)
DOCX = FINAL / "Asal_Glutamine_IFNG_Melanoma_Strengthened_Manuscript.docx"
TITLE = "Cellular Composition Accounts for Much of the Bulk Glutamine–IFN-γ Transcriptional Axis in Melanoma: A Cross-Cohort and Single-Cell Analysis"


def rows(path):
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def fmt_p(x):
    if x in (None, "", "NA"):
        return "—"
    v = float(x)
    return f"{v:.2e}" if v < 0.001 else f"{v:.3f}"


def fmt_n(x, digits=3):
    if x in (None, "", "NA"):
        return "—"
    return f"{float(x):.{digits}f}"


def set_font(run, size=11, bold=None, italic=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def keep(paragraph, lines=False):
    ppr = paragraph._p.get_or_add_pPr()
    node = OxmlElement("w:keepLines" if lines else "w:keepNext")
    ppr.append(node)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def shade(cell, fill="D9E8F5"):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def margins(cell):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for tag, value in (("top", 55), ("start", 65), ("bottom", 55), ("end", 65)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tcpr.append(tc_mar)


doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Inches(8.5), Inches(11)
section.top_margin, section.bottom_margin = Inches(0.82), Inches(0.78)
section.left_margin, section.right_margin = Inches(0.88), Inches(0.88)
section.header_distance, section.footer_distance = Inches(0.3), Inches(0.3)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.35
normal.paragraph_format.space_after = Pt(5)

for level, size in ((1, 15), (2, 12.5), (3, 11.5)):
    style = doc.styles[f"Heading {level}"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(11 if level == 1 else 7)
    style.paragraph_format.space_after = Pt(4)

for name, size in (("Figure Caption", 9), ("Table Caption", 9.5), ("Table Note", 8.5)):
    if name not in doc.styles:
        doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style = doc.styles[name]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = 1
    style.paragraph_format.space_after = Pt(5)

header = section.header.paragraphs[0]
header.text = "Asal | Cellular composition of the glutamine–IFN-γ axis"
for run in header.runs:
    set_font(run, size=8.5, color="666666")
page_number(section.footer.paragraphs[0])


def heading(text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    keep(paragraph)
    return paragraph


def para(text, lead=None):
    paragraph = doc.add_paragraph()
    if lead and text.startswith(lead):
        paragraph.add_run(lead).bold = True
        paragraph.add_run(text[len(lead):])
    else:
        paragraph.add_run(text)
    return paragraph


def bullet(text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    return paragraph


def table_caption(number, title):
    paragraph = doc.add_paragraph(style="Table Caption")
    paragraph.add_run(f"Table {number}. ").bold = True
    paragraph.add_run(title)
    keep(paragraph)


def table_note(text):
    paragraph = doc.add_paragraph(style="Table Note")
    paragraph.add_run("Note: ").italic = True
    paragraph.add_run(text)


def add_table(headers, body, size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, value in enumerate(headers):
        table.rows[0].cells[j].text = str(value)
    for values in body:
        cells = table.add_row().cells
        for j, value in enumerate(values):
            cells[j].text = str(value)
    for i, row in enumerate(table.rows):
        if i == 0:
            trpr = row._tr.get_or_add_trPr()
            rep = OxmlElement("w:tblHeader")
            rep.set(qn("w:val"), "true")
            trpr.append(rep)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            if i == 0:
                shade(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_font(run, size=size, bold=(i == 0))
    return table


def add_figure(number, path, caption, width=6.35, supplement=False):
    doc.add_page_break()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    label = f"Figure {'S' if supplement else ''}{number}"
    shape._inline.docPr.set("title", label)
    shape._inline.docPr.set("descr", caption)
    keep(paragraph)
    cap = doc.add_paragraph(style="Figure Caption")
    cap.add_run(f"{label}. ").bold = True
    cap.add_run(caption)
    keep(cap, lines=True)


# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(48)
set_font(p.add_run(TITLE), size=19, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
set_font(p.add_run("Islam Asal"), size=13, bold=True)
for line in ("OncoMetrika, Cairo, Egypt", "ORCID: 0009-0004-3187-7945", "Correspondence: contact through the repository"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    set_font(p.add_run(line), size=10.5)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(23)
p.add_run("Article type: ").bold = True
p.add_run("Original research / transcriptomic biomarker evaluation")
p.add_run("\nRunning title: ").bold = True
p.add_run("Cellular composition of the glutamine–IFN-γ axis")
p.add_run("\nKeywords: ").bold = True
p.add_run("melanoma; glutamine metabolism; interferon gamma; single-cell RNA sequencing; tumor microenvironment; survival; immunotherapy")
p.add_run("\nPreprint: ").bold = True
p.add_run("A preliminary version was posted on bioRxiv (doi: 10.1101/2025.09.28.679008). This manuscript reports a corrected and substantially expanded analysis.")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
p.add_run("Central finding").bold = True
p.add_run("\nThe inverse glutamine–IFN-γ relationship in bulk melanoma RNA is largely explained by cellular composition. IFN-γ-associated transcription is independently prognostic in two cohorts, whereas glutamine-associated transcription and the cross-score interaction are not consistently informative.")
doc.add_page_break()

heading("Abstract")
abstract = [
    ("Background", "An inverse relationship between glutamine-associated and interferon-γ (IFN-γ)-associated transcription in bulk melanoma could indicate metabolic antagonism or simply different cellular origins. We tested these alternatives and evaluated prognostic and immunotherapy-response associations."),
    ("Methods", "Continuous GSVA scores were analyzed in TCGA-SKCM (469 tumors), an independent melanoma survival cohort (GSE65904; 214 tumors), a pretreatment nivolumab cohort (GSE91061; 49 tumors), and single-cell RNA sequencing from 19 patients (GSE72056; 4,645 cells). Cox models used harmonized endpoints and clinical covariates. Bulk composition was examined with ABSOLUTE purity, EPIC, and MCP-counter. Single-cell scores were summarized by patient and compartment; paired tests treated patients, not cells, as independent units. Response models used Firth logistic regression, leave-one-out discrimination, and simulation-based precision analysis."),
    ("Results", "In TCGA, the candidate scores were inversely correlated before adjustment (r=−0.397; 95% CI −0.470 to −0.318) but progressively attenuated after purity/sample-type adjustment (r=−0.218), EPIC adjustment (r=−0.146), and MCP-counter adjustment (r=0.019). Single-cell analysis localized higher glutamine-associated scores to malignant cells than paired T, NK, or B-cell compartments, while IFN-γ-associated scores were lower in malignant cells than paired T-cell and macrophage compartments. IFN-γ was associated with lower mortality in TCGA (adjusted hazard ratio [HR] per SD 0.648; 95% CI 0.555–0.757) and GSE65904 (HR 0.667; 95% CI 0.534–0.835). Glutamine-associated scores were not independently prognostic in TCGA (HR 0.966; p=0.673) and were uncertain in GSE65904 (HR 0.823; p=0.084). The interaction was null in TCGA, nominal for the candidate sets in GSE65904, and null with curated sets. No response association was supported in GSE91061; approximately an odds ratio of 3.25 per SD was required for 80% simulated power."),
    ("Conclusions", "Cellular composition accounts for much of the bulk glutamine–IFN-γ transcriptional axis in melanoma. IFN-γ-associated transcription is a reproducible prognostic correlate, but glutamine-associated transcription provides no stable incremental outcome information. These data do not establish metabolic competition or treatment prediction; spatial or functional measurements and larger treatment-comparative cohorts are needed."),
]
for label, text in abstract:
    p = doc.add_paragraph()
    p.add_run(label + ": ").bold = True
    p.add_run(text)

heading("Introduction")
for text in [
    "Melanoma cells can depend on glutamine-derived carbon and nitrogen, while activated lymphocytes require coordinated nutrient availability to sustain effector function. Experimental studies therefore make metabolic competition within the tumor microenvironment biologically plausible [1–3]. IFN-γ-responsive transcription marks antigen presentation, chemokine signaling, cytotoxic inflammation, and adaptive immune resistance, and related expression profiles have been associated with benefit from PD-1 blockade [4].",
    "Bulk tumor transcriptomes, however, combine malignant, immune, stromal, and vascular signals. Immune-enriched transcription is associated with melanoma outcome [5], and single-cell studies show marked interpatient and intratumoral heterogeneity [9]. A negative bulk correlation between glutamine-associated and IFN-γ-associated scores could therefore arise from a biological interaction within cells, from variable proportions of cell populations that express the programs differently, or from both. Bulk correlation alone cannot distinguish these explanations.",
    "Our preliminary analysis proposed a median-defined four-group classification. Subsequent methodological review identified limitations in endpoint harmonization, small-sample response analysis, incremental-value assessment, and treatment of composition. More importantly, the original bulk-only analysis could not localize the two programs to cellular compartments.",
    "We therefore performed a strengthened, cross-cohort analysis with four complementary datasets. We asked whether (i) the bulk inverse association persisted after several composition adjustments, (ii) malignant and immune compartments showed different score localization at single-cell resolution, (iii) prognostic associations replicated in an independent survival cohort and with curated signatures, and (iv) pretreatment scores showed sufficiently precise evidence of immunotherapy-response association. The primary estimands used continuous scores; median groups were descriptive only."
]:
    para(text)

heading("Methods")
heading("Study design and data sources", 2)
para("This retrospective secondary analysis used public, deidentified transcriptomic and clinical data. Four cohorts served distinct roles: TCGA-SKCM for discovery survival and deconvolution; GSE65904 for independent survival validation; GSE91061 for exploratory pretreatment response analysis; and GSE72056 for single-cell localization. Because the treated dataset lacks a non-immunotherapy comparator, response associations are prognostic-within-treated-cohort associations and cannot establish treatment-predictive utility.")

heading("Bulk melanoma cohorts", 2)
para("TCGA-SKCM RNA-sequencing and clinical data were processed as described in the corrected reanalysis [5–7]. One expression profile per patient was retained, giving 469 patients. Overall survival (OS) time, event status, age, sex, and AJCC stage were obtained from the TCGA Clinical Data Resource [7]. The primary multivariable set contained 417 patients and 194 deaths. Stage was grouped as 0/I/II versus III/IV. ABSOLUTE purity and sample type were included in sensitivity analyses [12].")
para("GSE65904 is an independent Illumina HumanHT-12 v4 melanoma expression cohort [8]. Probe-level expression was mapped through GPL10558; multiple probes mapping to one gene were averaged. Disease-specific survival (DSS), age, sex, stage, and tissue source were harmonized before modeling. Of 214 available tumors, 210 had eligible DSS data and 203 had complete model covariates, including 99 disease-specific deaths.")
para("GSE91061 contains pretreatment melanoma biopsies from patients receiving nivolumab [6]. Complete or partial response was classified as response; stable or progressive disease as non-response; unknown response was excluded. Forty-nine pretreatment samples remained, including 10 responders. The analysis does not compare nivolumab with another treatment.")

heading("Single-cell cohort and annotation", 2)
para("GSE72056 contains single-cell RNA sequencing from 4,645 cells across 19 patients with metastatic melanoma [9]. We used the authors’ malignant-cell indicator and nonmalignant annotations. Cells with unresolved compartment identity were excluded from localization analyses, leaving 4,097 annotated cells: 1,257 malignant cells, 2,040 T cells, 512 B cells, 119 macrophages, 62 endothelial cells, 56 cancer-associated fibroblasts, and 51 NK cells. Analyses retained patient identity throughout.")

heading("Gene sets and continuous scores", 2)
para("The candidate glutamine-associated set contained 22 genes spanning glutamine handling and connected nitrogen/proline pathways; the candidate IFN-γ set contained 23 immune-response genes adapted from published IFN-γ-related profiles [1,4]. Bulk expression was log2 transformed where needed, and sample scores were calculated with GSVA and standardized within cohort [10]. Sensitivity analyses used MSigDB REACTOME_GLUTAMATE_AND_GLUTAMINE_METABOLISM and HALLMARK_INTERFERON_GAMMA_RESPONSE [11]. Gene mapping was documented for every cohort.")

heading("Bulk composition analyses", 2)
para("We first estimated the Pearson correlation between continuous candidate scores in TCGA. We then residualized each score for (i) ABSOLUTE purity and sample type, (ii) purity, sample type, and EPIC-derived immune/stromal fractions, and (iii) purity, sample type, and principal components of MCP-counter immune/stromal abundance scores [12–14]. EPIC fits with nonconvergence codes were excluded from EPIC-based models (20 of 469 samples). Correlations between each transcriptional score and deconvolution summaries were also calculated. EPIC- and MCP-counter-adjusted Cox models tested whether survival estimates persisted after composition adjustment.")

heading("Single-cell localization", 2)
para("Within GSE72056, each mapped signature gene was standardized across cells; the mean standardized expression of mapped genes defined a cell-level score. Scores were then summarized as the median within each patient-by-compartment combination. Malignant-cell scores were compared with each paired nonmalignant compartment using two-sided paired Wilcoxon tests with Holm adjustment across compartments. Median paired differences were accompanied by bootstrap 95% confidence intervals. Patient-level Spearman correlations assessed malignant glutamine-associated scores against malignant or aggregated immune IFN-γ-associated scores. The patient—not the cell—was the independent unit, avoiding pseudoreplication [20].")

heading("Outcome analyses", 2)
para("Cox proportional-hazards models used one-standard-deviation increments. TCGA models adjusted for age, sex, and grouped stage; GSE65904 models adjusted for age, sex, stage, and tissue source. Candidate models included both score main effects and their product interaction. Nested likelihood-ratio tests evaluated whether IFN-γ, glutamine, and the interaction added fit sequentially to clinical covariates. Proportional-hazards assumptions were assessed with scaled Schoenfeld residuals. Curated signatures and composition-adjusted models were prespecified sensitivity analyses.")
para("GSE91061 response associations used Firth penalized logistic regression because only 10 responses were observed [17]. Leave-one-out cross-validated predictions were summarized by ROC AUC with bootstrap confidence intervals [18]. Monte Carlo simulations used the observed sample size and response prevalence to estimate power across true odds ratios; this analysis quantifies the large effects detectable by the dataset rather than retroactively proving equivalence.")

heading("Missingness, multiplicity, and software", 2)
para("We reported cohort flow and compared score distributions and observed covariates between included and excluded cases. Primary hypotheses were estimated with continuous scores. Single-cell compartment comparisons used Holm adjustment; other sensitivity analyses are identified as such and interpreted by effect magnitude, confidence intervals, and consistency rather than isolated p-value thresholds. Analyses used R 4.4.2 with survival, GSVA, GEOquery, EPIC, MCPcounter, logistf, pROC, and related packages. Exact versions are recorded in the supplied session-information files.")

heading("Results")
heading("Cohorts and score mapping", 2)
para("The four datasets contributed complementary evidence (Table 1 and Figure 1). Candidate signatures mapped nearly completely in every dataset: 22 glutamine and 22 IFN-γ genes in TCGA, 21 and 23 in GSE65904, and 21 and 23 in the single-cell dataset. Curated sets mapped less completely for the small glutamine set but broadly for the 200-gene IFN-γ set. Cohort-specific mapping is reported in Supplementary Tables S2 and S5.")

heading("The bulk inverse association is strongly composition dependent", 2)
para("In TCGA, candidate glutamine-associated and IFN-γ-associated scores were inversely correlated (r=−0.397; 95% CI −0.470 to −0.318; p=3.89×10⁻¹⁹; Figure 2). Adjustment for purity and sample type attenuated the association to r=−0.218. Adding EPIC immune/stromal fractions attenuated it further to r=−0.146, whereas adjustment using MCP-counter immune/stromal principal components produced essentially no residual association (r=0.019; 95% CI −0.072 to 0.110; p=0.681; Table 3 and Figure 4).")
para("The candidate glutamine score was negatively correlated with inferred immune abundance by both EPIC (r=−0.448) and MCP-counter (r=−0.481), while the IFN-γ score was positively correlated with immune abundance (r=0.520 and r=0.800, respectively). This opposing compartment loading explains why the unadjusted bulk scores appear strongly antagonistic. The difference between EPIC and MCP-counter residual estimates also shows that the exact adjusted value depends on the composition model; the robust conclusion is substantial attenuation, not a single corrected correlation.")

heading("Single-cell analysis localizes the programs to different compartments", 2)
para("Patient-level single-cell summaries supported a compartmental explanation (Figure 5). Candidate glutamine-associated scores were higher in malignant cells than paired T cells (median difference 0.424; Holm p=0.010), NK cells (0.398; p=0.037), or B cells (0.454; p=0.019). Candidate IFN-γ-associated scores were lower in malignant cells than paired T cells (median difference −0.356; Holm p=0.0066) and macrophages (−0.806; p=0.026). Curated signatures yielded directionally similar localization, although significance varied with mapping and the small number of paired patients.")
para("Across 15 evaluable patients, the correlation between malignant-cell candidate glutamine scores and aggregated immune-cell candidate IFN-γ scores was imprecise (Spearman ρ=−0.386; bootstrap 95% CI −0.842 to 0.256; p=0.156). Thus, cell-type localization is supported, but this small dataset does not establish a patient-level antagonistic mechanism.")

heading("IFN-γ prognosis replicates; glutamine prognosis does not", 2)
para("In the TCGA primary multivariable model (n=417; 194 deaths), higher IFN-γ-associated score was associated with lower mortality (HR per SD 0.648; 95% CI 0.555–0.757; p=4.89×10⁻⁸), whereas the glutamine-associated score was not (HR 0.966; 95% CI 0.822–1.135; p=0.673). Adding IFN-γ to the clinical model improved fit (likelihood-ratio p=8.20×10⁻⁹; concordance 0.636 to 0.691). Adding glutamine after IFN-γ did not improve fit (p=0.616).")
para("The IFN-γ association replicated in GSE65904 (n=203; 99 DSS deaths; adjusted HR 0.667; 95% CI 0.534–0.835; p=3.90×10⁻⁴). The glutamine estimate was directionally protective but uncertain (HR 0.823; 95% CI 0.660–1.026; p=0.084). IFN-γ also improved the external clinical model (likelihood-ratio p=0.000505; concordance 0.621 to 0.676). Curated signatures confirmed the IFN-γ result in both cohorts and did not support an independent glutamine association. Figure 3 summarizes cross-cohort estimates.")
para("Composition-adjusted TCGA models did not materially change the IFN-γ result: HR 0.668 (95% CI 0.550–0.811; p=4.40×10⁻⁵) after EPIC adjustment and HR 0.774 (95% CI 0.596–1.004; p=0.054) after MCP-counter adjustment. The latter estimate was less precise but directionally consistent. Glutamine remained null. Proportional-hazards tests did not identify a global violation in the external model.")

heading("The interaction is not a replicated biomarker", 2)
para("The candidate glutamine-by-IFN-γ interaction was null in TCGA (HR 1.027; 95% CI 0.882–1.196; p=0.728). In GSE65904, the candidate interaction was nominally associated with DSS (HR 1.252; 95% CI 1.001–1.566; p=0.049), and the nested interaction test gave p=0.044. However, the curated-signature interaction was null in GSE65904 (HR 1.085; p=0.445), and neither candidate nor curated interactions were supported in TCGA. We therefore classify the external candidate interaction as an unstable, exploratory observation rather than validation.")

heading("The treated cohort is too small for modest response effects", 2)
para("In GSE91061, neither candidate score was associated with response in Firth models: glutamine odds ratio (OR) 0.697 (95% CI 0.322–1.405; p=0.319) and IFN-γ OR 1.135 (95% CI 0.564–2.351; p=0.722). The exploratory interaction OR was 0.680 (95% CI 0.279–1.426; p=0.323). Leave-one-out AUC was 0.356 (95% bootstrap CI 0.156–0.556; Figure 6). All score intervals were compatible with both decreases and increases in response odds.")
para("At the observed prevalence of 10 responses among 49 patients, simulation indicated that an OR of approximately 3.25 per SD was required to reach 80% power under the modeled conditions (Supplementary Figure S3). Accordingly, the null estimates do not exclude modest associations. They do show that this cohort cannot credibly validate a multivariable predictive signature.")

heading("Missingness and robustness", 2)
para("TCGA patients included in the complete-covariate model had score distributions similar to patients with expression but incomplete outcome/covariate data. Stage distribution differed between OS-eligible and expression-only groups (p=0.035), and stage missingness drove most primary-model exclusions. Included patients had shorter observed OS time than OS-eligible excluded patients, while score distributions remained similar. These patterns limit transportability and reinforce interpretation of the models as adjusted associations, not population-level risk tools.")

heading("Discussion")
heading("Principal findings", 2)
para("This strengthened analysis changes the biological interpretation of the original observation. The inverse relationship between glutamine-associated and IFN-γ-associated transcription is real in bulk melanoma RNA, but it is not robust to modeling cellular composition. Independent deconvolution methods and single-cell localization converge on a parsimonious explanation: glutamine-associated transcription is relatively enriched in malignant cells, whereas IFN-γ-associated transcription is concentrated in immune compartments. Variation in the mixture of these compartments generates much of the apparent bulk antagonism.")
para("The most reproducible outcome result is not a two-score metabolic-conflict biomarker but the favorable prognostic association of IFN-γ-associated transcription. This association replicated across OS and DSS cohorts, survived curated-signature analyses, and remained directionally consistent after composition adjustment. Glutamine-associated transcription neither consistently predicted survival nor improved models after IFN-γ. The nominal external interaction did not reproduce across signatures or cohorts and should not motivate clinical stratification without new data.")

heading("Biological interpretation", 2)
para("The findings do not refute metabolic competition. Experimental work shows that nutrient restriction can alter tumor and immune-cell function [2,3]. Rather, they demonstrate that bulk coexpression is insufficient evidence for that mechanism. A score can reflect cellular abundance, a cell-intrinsic state, or both. Here, the near-null MCP-counter-adjusted correlation and the single-cell compartment differences place cellular mixture at the center of the bulk signal.")
para("The remaining uncertainty is mechanistic and spatial. The single-cell dataset had only 19 patients, lacked matched functional metabolite measurements, and used dissociated cells rather than spatially resolved neighborhoods. Patient-level cross-compartment correlations had wide confidence intervals. Future work should pair spatial transcriptomics or multiplex imaging with isotope tracing, nutrient measurements, or perturbation assays to determine whether glutamine-active malignant neighborhoods suppress nearby IFN-γ programs.")

heading("Clinical interpretation", 2)
para("IFN-γ-associated transcription is prognostic in these datasets, but prognostic association does not establish immunotherapy prediction. Demonstrating predictive value requires treatment-by-biomarker interaction in a randomized or otherwise credible treatment-comparative design. The single treated cohort analyzed here is small, includes only 10 responders, and can detect only very large effects. The present data therefore provide no basis for treatment selection or clinical cut points.")

heading("Strengths and limitations", 2)
para("Strengths include independent survival validation, continuous-score modeling, curated-signature sensitivity analyses, three complementary composition approaches, patient-level single-cell inference, explicit incremental-value tests, small-sample response methods, and complete machine-readable outputs. These elements address the principal weaknesses of the preliminary version.")
para("Limitations remain. All analyses are retrospective. Cohorts differed in platform, tissue source, endpoint, and clinical annotation. Deconvolution relies on reference profiles and model assumptions; EPIC did not fully converge for 20 tumors, which were excluded from EPIC models. The single-cell data include few patients and uneven cell-type recovery. Signature scores are transcriptional proxies rather than glutamine flux or IFN-γ protein activity. Missing clinical covariates may induce selection bias. Multiple sensitivity analyses increase the chance of nominal findings, particularly the isolated GSE65904 interaction. Finally, no treatment-comparative cohort was available.")

heading("Conclusions")
para("Cellular composition accounts for much of the inverse bulk glutamine–IFN-γ transcriptional axis in melanoma. IFN-γ-associated transcription is a reproducible favorable prognostic correlate, but glutamine-associated transcription and the cross-score interaction do not provide stable incremental outcome information. The results support cell-resolved mechanistic studies and adequately powered treatment-comparative validation—not a median-defined two-score clinical classifier.")

heading("Declarations")
heading("Ethics approval and consent", 2)
para("Only public, deidentified data were analyzed; no new participant recruitment or intervention was performed. Ethical approvals and consent procedures are described in the source studies.")
heading("Data availability", 2)
para("TCGA-SKCM data are available through the NCI Genomic Data Commons. GEO datasets are available under GSE65904, GSE91061, and GSE72056. The preliminary preprint is available at doi: 10.1101/2025.09.28.679008.")
heading("Code availability", 2)
para("Complete analysis scripts, derived tables, figures, data-acquisition manifests, and software session information accompany this revision package. A permanent public repository URL and archived release identifier should be inserted before journal submission.")
heading("Competing interests", 2)
para("The author declares no competing interests.")
heading("Funding", 2)
para("No external funding was received for this work.")
heading("Author contributions", 2)
para("Islam Asal: conceptualization, methodology, software, formal analysis, visualization, writing—original draft, and writing—review and editing.")
heading("Acknowledgments", 2)
para("The author thanks the investigators and participants who generated and shared the TCGA and GEO datasets, and the PREreview contributors whose methodological critique motivated the strengthened analysis.")

heading("References")
references = [
    "1. Ratnikov B, Aza-Blanc P, Ronai ZA, Smith JW, Osterman AL, Scott DA. Glutamate and asparagine cataplerosis underlie glutamine addiction in melanoma. Oncotarget. 2015;6:7379–7389. doi:10.18632/oncotarget.3132.",
    "2. Chang CH, Qiu J, O’Sullivan D, et al. Metabolic competition in the tumor microenvironment is a driver of cancer progression. Cell. 2015;162:1229–1241. doi:10.1016/j.cell.2015.08.016.",
    "3. Leone RD, Zhao L, Englert JM, et al. Glutamine blockade induces divergent metabolic programs to overcome tumor immune evasion. Science. 2019;366:1013–1021. doi:10.1126/science.aav2588.",
    "4. Ayers M, Lunceford J, Nebozhyn M, et al. IFN-γ-related mRNA profile predicts clinical response to PD-1 blockade. J Clin Invest. 2017;127:2930–2940. doi:10.1172/JCI91190.",
    "5. Cancer Genome Atlas Network. Genomic classification of cutaneous melanoma. Cell. 2015;161:1681–1696. doi:10.1016/j.cell.2015.05.044.",
    "6. Riaz N, Havel JJ, Makarov V, et al. Tumor and microenvironment evolution during immunotherapy with nivolumab. Cell. 2017;171:934–949.e16. doi:10.1016/j.cell.2017.09.028.",
    "7. Liu J, Lichtenberg T, Hoadley KA, et al. An integrated TCGA pan-cancer clinical data resource to drive high-quality survival outcome analytics. Cell. 2018;173:400–416.e11. doi:10.1016/j.cell.2018.02.052.",
    "8. Cirenajwis H, Ekedahl H, Lauss M, et al. Molecular stratification of metastatic melanoma using gene expression profiling: prediction of survival outcome and benefit from molecular targeted therapy. Oncotarget. 2015;6:12297–12309. doi:10.18632/oncotarget.3655.",
    "9. Tirosh I, Izar B, Prakadan SM, et al. Dissecting the multicellular ecosystem of metastatic melanoma by single-cell RNA-seq. Science. 2016;352:189–196. doi:10.1126/science.aad0501.",
    "10. Hänzelmann S, Castelo R, Guinney J. GSVA: gene set variation analysis for microarray and RNA-seq data. BMC Bioinformatics. 2013;14:7. doi:10.1186/1471-2105-14-7.",
    "11. Liberzon A, Birger C, Thorvaldsdóttir H, et al. The Molecular Signatures Database Hallmark Gene Set Collection. Cell Syst. 2015;1:417–425. doi:10.1016/j.cels.2015.12.004.",
    "12. Carter SL, Cibulskis K, Helman E, et al. Absolute quantification of somatic DNA alterations in human cancer. Nat Biotechnol. 2012;30:413–421. doi:10.1038/nbt.2203.",
    "13. Becht E, Giraldo NA, Lacroix L, et al. Estimating the population abundance of tissue-infiltrating immune and stromal cell populations using gene expression. Genome Biol. 2016;17:218. doi:10.1186/s13059-016-1070-5.",
    "14. Racle J, de Jonge K, Baumgaertner P, Speiser DE, Gfeller D. Simultaneous enumeration of cancer and immune cell types from bulk tumor gene expression data. eLife. 2017;6:e26476. doi:10.7554/eLife.26476.",
    "15. Davis S, Meltzer PS. GEOquery: a bridge between the Gene Expression Omnibus and BioConductor. Bioinformatics. 2007;23:1846–1847. doi:10.1093/bioinformatics/btm254.",
    "16. Colaprico A, Silva TC, Olsen C, et al. TCGAbiolinks: an R/Bioconductor package for integrative analysis of TCGA data. Nucleic Acids Res. 2016;44:e71. doi:10.1093/nar/gkv1507.",
    "17. Heinze G, Schemper M. A solution to the problem of separation in logistic regression. Stat Med. 2002;21:2409–2419. doi:10.1002/sim.1047.",
    "18. Robin X, Turck N, Hainard A, et al. pROC: an open-source package for R and S+ to analyze and compare ROC curves. BMC Bioinformatics. 2011;12:77. doi:10.1186/1471-2105-12-77.",
    "19. McShane LM, Altman DG, Sauerbrei W, et al. REporting recommendations for tumor MARKer prognostic studies (REMARK). J Natl Cancer Inst. 2005;97:1180–1184. doi:10.1093/jnci/dji237.",
    "20. Squair JW, Gautier M, Kathe C, et al. Confronting false discoveries in single-cell differential expression. Nat Commun. 2021;12:5692. doi:10.1038/s41467-021-25960-2.",
    "21. Sade-Feldman M, Yizhak K, Bjorgaard SL, et al. Defining T cell states associated with response to checkpoint immunotherapy in melanoma. Cell. 2018;175:998–1013.e20. doi:10.1016/j.cell.2018.10.038."
]
for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = Pt(4)
    p.add_run(ref)

# Main tables
doc.add_page_break()
heading("Tables")
table_caption(1, "Cohort roles and primary analysis sets")
cohorts = rows(NEW / "tables" / "Table1_strengthened_cohort_flow.csv")
add_table(["Cohort", "Role", "Available", "Primary analysis", "Events/responders"],
          [[r["cohort"], r["role"], r["available"], r["primary_analysis"], r["events_or_responders"]] for r in cohorts], size=8)
table_note("For GSE72056, the primary unit is the patient; 4,645 cells from 19 patients were available and 4,097 annotated cells contributed to localization analyses.")

doc.add_page_break()
table_caption(2, "Cross-cohort multivariable survival associations for continuous scores")
tcga = rows(OLD / "tables" / "Table2_cox_models.csv")
gse = rows(NEW / "tables" / "Table5_GSE65904_cox_models.csv")
body = []
for r in tcga:
    if r["model"] == "Clinical covariates" and r["term"] in {"gln_z", "ifng_z", "gln_z:ifng_z"}:
        term = {"gln_z":"Glutamine", "ifng_z":"IFN-γ", "gln_z:ifng_z":"Interaction"}[r["term"]]
        body.append(["TCGA-SKCM", "Candidate", term, r["n"], r["events"], fmt_n(r["estimate"]), f"{fmt_n(r['conf.low'])}–{fmt_n(r['conf.high'])}", fmt_p(r["p.value"])])
for r in gse:
    if r["model"] in {"Candidate signatures + clinical covariates", "Curated signatures + clinical covariates"} and r["term"] in {"gln_z", "ifng_z", "gln_z:ifng_z", "curated_gln_z", "curated_ifng_z", "curated_gln_z:curated_ifng_z"}:
        definition = "Curated" if "Curated" in r["model"] else "Candidate"
        term = "Interaction" if ":" in r["term"] else ("IFN-γ" if "ifng" in r["term"] else "Glutamine")
        body.append(["GSE65904", definition, term, r["n"], r["events"], fmt_n(r["hazard_ratio"]), f"{fmt_n(r['conf.low'])}–{fmt_n(r['conf.high'])}", fmt_p(r["p.value"])])
add_table(["Cohort", "Definition", "Term", "n", "Events", "HR/SD", "95% CI", "p"], body, size=7.7)
table_note("TCGA models adjusted for age, sex, and grouped stage. GSE65904 models adjusted for age, sex, stage, and tissue source. HR, hazard ratio.")

doc.add_page_break()
table_caption(3, "Attenuation of the TCGA bulk score correlation after composition adjustment")
cor = rows(NEW / "tables" / "Table6_bulk_composition_partial_correlations.csv")
add_table(["Adjustment", "n", "Correlation", "95% CI", "p"],
          [[r["adjustment"], r["n"], fmt_n(r["correlation"]), f"{fmt_n(r['conf.low'])}–{fmt_n(r['conf.high'])}", fmt_p(r["p.value"])] for r in cor], size=8)
table_note("Pearson correlations of continuous candidate signature scores. Adjusted values are correlations between model residuals. Twenty EPIC fits with nonconvergence codes were excluded from EPIC analyses.")

doc.add_page_break()
table_caption(4, "Patient-level single-cell localization of candidate signatures")
loc = rows(NEW / "tables" / "Table9_GSE72056_paired_localization.csv")
loc = [r for r in loc if r["signature"] in {"gln_score", "ifng_score"}]
add_table(["Signature", "Comparison", "Paired patients", "Median difference", "95% CI", "Raw p", "Holm p"],
          [["Glutamine" if r["signature"] == "gln_score" else "IFN-γ", r["comparison"], r["paired_patients"], fmt_n(r["median_difference"]), f"{fmt_n(r['conf.low'])}–{fmt_n(r['conf.high'])}", fmt_p(r["p.value"]), fmt_p(r["p.adjusted"])] for r in loc], size=7.4)
table_note("Positive differences indicate higher scores in malignant cells. Scores were summarized per patient and compartment before paired testing; cells were not treated as independent replicates.")

doc.add_page_break()
table_caption(5, "Exploratory pretreatment response models in GSE91061")
resp = rows(OLD / "tables" / "Table3_firth_logistic_models.csv")
resp = [r for r in resp if r["term"] in {"gln_z", "ifng_z", "gln_z:ifng_z"} and "Candidate" in r["model"]]
add_table(["Model", "Term", "OR/SD", "95% CI", "p"],
          [[r["model"].replace("Candidate signatures - ", ""), {"gln_z":"Glutamine", "ifng_z":"IFN-γ", "gln_z:ifng_z":"Interaction"}[r["term"]], fmt_n(r["odds_ratio"]), f"{fmt_n(r['conf.low'])}–{fmt_n(r['conf.high'])}", fmt_p(r["p.value"])] for r in resp], size=8)
table_note("Firth penalized logistic regression; n=49 with 10 responders. OR, odds ratio. AUC values were calculated from leave-one-out predictions.")

# Main figures
add_figure(1, OLD / "figures" / "Figure1_TCGA_cohort_flow.png", "TCGA-SKCM analysis flow from aligned expression profiles to the overall-survival, complete-covariate, and composition-adjusted analysis sets. Table 1 reports the complementary external cohorts.")
add_figure(2, OLD / "figures" / "Figure2_TCGA_score_correlation.png", "Inverse association between continuous candidate glutamine-associated and IFN-γ-associated GSVA scores in TCGA-SKCM. The line is the least-squares fit; the shaded band is its 95% confidence interval. This unadjusted bulk association is not interpreted as cell-intrinsic antagonism.")
add_figure(3, NEW / "figures" / "Figure5_cross_cohort_survival_forest.png", "Cross-cohort adjusted survival associations per one-standard-deviation increase in score. IFN-γ-associated transcription is directionally consistent and statistically supported in TCGA OS and GSE65904 DSS. Glutamine and interaction estimates are not consistently supported. The dashed line marks HR=1.")
add_figure(4, NEW / "figures" / "Figure6_bulk_cell_composition.png", "Progressive attenuation of the TCGA bulk glutamine–IFN-γ score correlation after adjustment for tumor purity, sample type, and immune/stromal composition. Points show Pearson correlations of residualized continuous scores; bars show 95% confidence intervals.")
add_figure(5, NEW / "figures" / "Figure7_single_cell_localization.png", "Patient-level single-cell localization in GSE72056. Each point is a patient-specific median score for one annotated compartment; boxplots summarize patients, not cells. Candidate glutamine-associated scores are relatively enriched in malignant cells, whereas IFN-γ-associated scores are enriched in immune compartments.")
add_figure(6, OLD / "figures" / "Figure4_GEO_response.png", "Exploratory pretreatment score distributions by response in GSE91061. The cohort contained 49 tumors and 10 responders. Firth regression and leave-one-out discrimination did not support a response association, but confidence intervals were wide.")

# Supplement
doc.add_page_break()
heading("Supplementary material")
heading("Supplementary methods", 2)
para("Data acquisition was scripted and cached with MD5 checksums. GSE65904 series-matrix and GPL10558 annotation files were retrieved from GEO; GSE72056 revised single-cell data were retrieved from GEO; MCP-counter gene markers were pinned to repository commit b6eac73. The acquisition manifest records local filenames, source URLs, sizes, and checksums. All strengthened analyses can be regenerated by running 05_acquire_strengthening_data.R followed by 04_revised_analysis.R and 06_strengthened_analysis.R, subject to availability of the public source files and required R packages.")
heading("Supplementary results", 2)
para("Curated-signature analyses reproduced the favorable IFN-γ survival association and did not support a stable glutamine or interaction effect. In GSE65904, curated IFN-γ HR was 0.735 (95% CI 0.591–0.915; p=0.0059), curated glutamine HR was 0.998 (p=0.988), and the interaction HR was 1.085 (p=0.445). In TCGA, corresponding estimates were 0.636 (p=1.44×10⁻⁸), 0.947 (p=0.507), and 0.992 (p=0.915).")
para("After EPIC adjustment in TCGA, IFN-γ remained associated with survival (HR 0.668; p=4.40×10⁻⁵), glutamine remained null (HR 0.939; p=0.495), and the interaction remained null (HR 0.963; p=0.667). MCP-counter-adjusted estimates were less precise: IFN-γ HR 0.774 (p=0.054), glutamine HR 0.917 (p=0.334), and interaction HR 0.998 (p=0.984).")

doc.add_page_break()
table_caption("S1", "Incremental model fit in the independent GSE65904 survival cohort")
inc = rows(NEW / "tables" / "TableS6_GSE65904_incremental_models.csv")
add_table(["Model", "n", "Events", "Parameters", "AIC", "Concordance", "Incremental LRT p"],
          [[r["model"], r["n"], r["events"], r["parameters"], fmt_n(r["AIC"], 2), fmt_n(r["concordance"]), fmt_p(r["LRT_vs_previous"])] for r in inc], size=7.5)
table_note("Nested likelihood-ratio tests compare each row with the preceding model. Candidate signatures are shown.")

doc.add_page_break()
table_caption("S2", "Signature mapping across cohorts")
mapping = [("GSE65904", r) for r in rows(NEW / "tables" / "TableS8_GSE65904_signature_mapping.csv")] + [("GSE72056", r) for r in rows(NEW / "tables" / "TableS11_GSE72056_signature_mapping.csv")]
add_table(["Cohort", "Signature", "Genes specified", "Genes mapped"],
          [[cohort, r["signature"], r["defined_genes"], r["mapped_genes"]] for cohort, r in mapping], size=7.8)
table_note("The complete machine-readable mapping tables are supplied with the analysis package.")

doc.add_page_break()
table_caption("S3", "TCGA score and covariate missingness comparisons")
miss = rows(NEW / "tables" / "TableS12_TCGA_missingness_comparison.csv")
add_table(["Comparison", "Variable", "Included/eligible", "Excluded/other", "p"],
          [[r.get("comparison", ""), r.get("variable", ""), r.get("group1_summary", r.get("included_summary", "")), r.get("group2_summary", r.get("excluded_summary", "")), fmt_p(r.get("p.value", r.get("p_value", "NA")))] for r in miss], size=7)
table_note("Summaries and tests are exploratory diagnostics of selection into analytic sets.")

add_figure(1, OLD / "figures" / "Figure3_TCGA_Kaplan_Meier.png", "Descriptive TCGA Kaplan–Meier curves for median-defined score groups. These groups were not the primary estimand and should not be used as clinical cut points.", supplement=True)
add_figure(2, OLD / "figures" / "FigureS1_score_distributions.png", "Continuous score distributions in TCGA-SKCM and GSE91061.", supplement=True)
add_figure(3, NEW / "figures" / "FigureS5_GSE91061_power_curve.png", "Monte Carlo power under the observed GSE91061 sample size (n=49) and response prevalence (10/49). Approximately OR=3.25 per SD was required for 80% power under the simulation model.", supplement=True)

doc.core_properties.title = TITLE
doc.core_properties.author = "Islam Asal"
doc.core_properties.subject = "Cross-cohort and single-cell melanoma transcriptomic analysis"
doc.core_properties.keywords = "melanoma, glutamine, IFN-gamma, single-cell RNA sequencing, survival"
doc.save(DOCX)
print(DOCX)
